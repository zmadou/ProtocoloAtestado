#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Form → Planilha + PDF (multi-sigla, dark theme, progress bar, validações)
 - SIGLA configurável e persistente (MCI, MMD, IOT, etc.)
 - Cria planilha como MALA<SIGLA>.xlsx (ex.: MALAMMD.xlsx)
 - Gera DOCX/PDF em "Requerimentos/"
 - Mantém sequência do protocolo baseada **apenas** na planilha (coluna 'N req.')
 - Se (Nome+ID+CPF+Data) já existir na planilha, NÃO duplica — só (re)gera o arquivo.
 - Tema escuro com azul, barra de progresso verde
 - Validações/máscaras: CPF (somente números, 11 dígitos, formatado no DOC),
   ID tamanho fixo, Código da oferta tamanho fixo, Nº chamado numérico (livre).
"""

import os
import re
import json
from datetime import datetime
from typing import Dict, Any, Optional

import pandas as pd
from docx import Document

# ---------------- Config base ----------------
DEFAULT_ANO = os.environ.get("MCI_ANO", "2025")
DEFAULT_MODELO = os.environ.get("MCI_DOCX", "anexo_geduc_se_requerimento_amparo_legall.docx")
PASTA_SAIDA = os.environ.get("MCI_SAIDAS", "Requerimentos")
CONFIG_PATH = os.environ.get("MCI_CONFIG", "config_form.json")

# Defaults de validação (podem ser sobrescritos pelo config_form.json)
DEFAULTS_VALID = {"cpf_len": 11, "id_len": 10, "oferta_len": 10}

MESES = {1:"janeiro",2:"fevereiro",3:"março",4:"abril",5:"maio",6:"junho",
         7:"julho",8:"agosto",9:"setembro",10:"outubro",11:"novembro",12:"dezembro"}

COLUNAS_PADRAO = [
    "N req.","N chamado","NOME","ID","CPF","CURSO","TURMA",
    "Código da oferta","Data","retorno (Previsão)"
]

# ---------------- Utilidades de SIGLA/CONFIG ----------------
def _sanitize_sigla(s: str) -> str:
    s = (s or "").strip().upper()
    s = re.sub(r"[^A-Z0-9]", "", s)
    return s or "MCI"

def carregar_config() -> Dict[str, Any]:
    cfg = {
        "sigla": _sanitize_sigla(os.environ.get("MCI_SIGLA", "MCI")),
        "ano": str(DEFAULT_ANO),
        "valid": DEFAULTS_VALID.copy(),
        "last_req": 0
    }
    if os.path.isfile(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                if "sigla" in data:
                    cfg["sigla"] = _sanitize_sigla(data["sigla"])
                if "ano" in data and str(data["ano"]).isdigit():
                    cfg["ano"] = str(data["ano"])
                if "valid" in data and isinstance(data["valid"], dict):
                    cfg["valid"].update({k:int(v) for k,v in data["valid"].items() if str(v).isdigit()})
                if "last_req" in data and str(data["last_req"]).isdigit():
                    cfg["last_req"] = int(data["last_req"])
        except Exception:
            pass
    return cfg

def salvar_config(sigla: str, ano: Optional[str] = None,
                  valid: Optional[Dict[str,int]] = None,
                  last_req: Optional[int] = None) -> None:
    data = {"sigla": _sanitize_sigla(sigla), "ano": str(ano or DEFAULT_ANO)}
    if valid:
        data["valid"] = valid
    if last_req is not None:
        data["last_req"] = int(last_req)
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def nome_planilha(sigla: str) -> str:
    forced = os.environ.get("MCI_XLSX", "").strip()
    return forced if forced else f"MALA{_sanitize_sigla(sigla)}.xlsx"

# ---------------- Funções auxiliares ----------------
def fmt_data_pt(dt: datetime) -> str:
    return f"{dt.day} de {MESES[dt.month]} de {dt.year}"

def sanitize_num(s: Any) -> str:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    return re.sub(r"\D", "", str(s))

def get_str(s: Any) -> str:
    return "" if (s is None or (isinstance(s, float) and pd.isna(s))) else str(s).strip()

def parse_data_flex(valor: str, fallback_hoje: bool = True) -> datetime:
    if not valor:
        return datetime.today() if fallback_hoje else None
    for fmt in ("%d/%m/%Y","%Y-%m-%d","%d-%m-%Y","%d.%m.%Y"):
        try:
            return datetime.strptime(str(valor), fmt)
        except Exception:
            pass
    try:
        dt = pd.to_datetime(valor, dayfirst=True, errors="coerce")
        if pd.isna(dt):
            raise ValueError
        return dt.to_pydatetime()
    except Exception:
        return datetime.today() if fallback_hoje else None

def carregar_planilha(caminho: str) -> pd.DataFrame:
    if os.path.isfile(caminho):
        df = pd.read_excel(caminho)
        df.columns = [c.strip() for c in df.columns]
        return df
    return pd.DataFrame(columns=COLUNAS_PADRAO)

def salvar_planilha(df: pd.DataFrame, caminho: str) -> None:
    for c in COLUNAS_PADRAO:
        if c not in df.columns:
            df[c] = None
    df.to_excel(caminho, index=False)

def proximo_nreq(df: pd.DataFrame, last_req_cfg: int = 0) -> int:
    """Próximo número baseado **apenas** na planilha; se vazia, usa fallback do config."""
    if "N req." in df.columns and not df.empty:
        nums = pd.to_numeric(df["N req."], errors="coerce")
        if not nums.dropna().empty:
            return int(nums.max()) + 1
    return last_req_cfg + 1 if last_req_cfg else 1

def existe_registro(df: pd.DataFrame, dados: Dict[str, Any]) -> Optional[int]:
    if df.empty:
        return None
    filtros = (df.get("NOME","").astype(str).str.strip().str.casefold()
               == str(dados.get("NOME","")).strip().casefold())
    for col in ("ID","CPF","Data"):
        if col in df.columns:
            filtros = filtros & (df[col].astype(str).str.strip() == str(dados.get(col,"")).strip())
    existentes = df[filtros]
    if existentes.empty:
        return None
    try:
        return int(pd.to_numeric(existentes.iloc[0]["N req."], errors="coerce"))
    except Exception:
        return None

def format_cpf(digs: str) -> str:
    d = sanitize_num(digs)[:11]
    return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}" if len(d) == 11 else d

def montar_mapa(row: Dict[str, Any], protocolo_num: int) -> Dict[str, str]:
    dt_solic = parse_data_flex(row.get("Data"))
    dt_retorno = parse_data_flex(row.get("retorno (Previsão)"))
    return {
        "protocolo": f"{int(protocolo_num):02d}",
        "nome": get_str(row.get("NOME")),
        "id": sanitize_num(row.get("ID")),
        "cpf": format_cpf(row.get("CPF")),  # formatado só na saída
        "curso": get_str(row.get("CURSO")),
        "turma": get_str(row.get("TURMA")),
        "oferta": sanitize_num(row.get("Código da oferta")),
        "chamado": sanitize_num(row.get("N chamado")),
        "data_ext": fmt_data_pt(dt_solic),
        "data_retorno_ext": fmt_data_pt(dt_retorno),
    }

def substituir_texto(doc: Document, mapa: Dict[str, str], sigla: str, ano: str) -> None:
    sigla = _sanitize_sigla(sigla)
    pares = [
        (r"Protocolo nº \d{2}-[A-Z0-9]+/\d{4}", f"Protocolo nº {mapa['protocolo']}-{sigla}/{ano}"),
        (r"Eu, .+?, ID nº .*?; CPF nº .*?, estudante regularmente matriculado\(a\) no curso .*?,\s*TURMA:.*?,",
         f"Eu, {mapa['nome']}, ID nº {mapa['id']}; CPF nº {mapa['cpf']}, estudante regularmente matriculado(a) no curso {mapa['curso']}, TURMA:{mapa['turma']},"),
        (r"Código da oferta: .*? \(preenchimento do setor da secretaria escolar\)",
         f"Código da oferta: {mapa['oferta']} (preenchimento do setor da secretaria escolar)"),
        (r"São Paulo, .*?\.", f"São Paulo, {mapa['data_ext']}."),
        (r"Conforme chamado de nº .*", f"Conforme chamado de nº {mapa['chamado']}"),
        (r"Aluno .+", f"Aluno {mapa['nome']}"),
        (r"Data de retorno até: .*?\s+\(considerar.*\)",
         f"Data de retorno até: {mapa['data_retorno_ext']}  (considerar de 1 a 7 dias úteis, a partir da data de solicitação)"),
    ]
    def apply_replace(text: str) -> str:
        for pattern, repl in pares:
            text = re.sub(pattern, repl, text)
        return text

    for p in doc.paragraphs:
        if p.text:
            joined = "".join(run.text for run in p.runs)
            new = apply_replace(joined)
            if new != joined:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        joined = "".join(run.text for run in p.runs)
                        new = apply_replace(joined)
                        if new != joined:
                            p.runs[0].text = new
                            for r in p.runs[1:]:
                                r.text = ""

# ---------- Conversão DOCX → PDF (com encerramento limpo COM/Word) ----------
def tentar_converter_pdf(docx_path: str, pdf_path: str) -> bool:
    # 1) docx2pdf
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(docx_path, pdf_path)
        return os.path.isfile(pdf_path)
    except Exception:
        pass
    # 2) comtypes
    try:
        import comtypes.client
        import pythoncom
        pythoncom.CoInitialize()
        word = None
        try:
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close(False)
            return os.path.isfile(pdf_path)
        finally:
            if word is not None:
                try: word.Quit()
                except Exception: pass
            try: pythoncom.CoUninitialize()
            except Exception: pass
    except Exception:
        pass
    # 3) win32com
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close(False)
            return os.path.isfile(pdf_path)
        finally:
            if word is not None:
                try: word.Quit()
                except Exception: pass
            try: pythoncom.CoUninitialize()
            except Exception: pass
    except Exception:
        pass
    return False

def gerar_documento(linha: Dict[str, Any], protocolo_num: int, sigla: str, ano: str, modelo: str) -> str:
    os.makedirs(PASTA_SAIDA, exist_ok=True)
    mapa = montar_mapa(linha, protocolo_num)
    base = f"{mapa['protocolo']}{_sanitize_sigla(sigla)}{ano} {mapa['nome']}"
    out_docx = os.path.join(PASTA_SAIDA, base + ".docx")
    out_pdf = os.path.join(PASTA_SAIDA, base + ".pdf")

    if os.path.isfile(out_pdf):
        return out_pdf
    if os.path.isfile(out_docx):
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(out_docx, out_pdf)
            return out_pdf
        except Exception:
            return out_docx

    doc = Document(modelo)
    substituir_texto(doc, mapa, sigla, ano)
    doc.save(out_docx)

    saida_final = out_docx
    try:
        if tentar_converter_pdf(out_docx, out_pdf):
            saida_final = out_pdf
    except Exception:
        pass
    return saida_final

# ---------------- GUI ----------------
def iniciar_form():
    import tkinter as tk
    from tkinter import ttk, messagebox
    import atexit

    # Tema (mantendo seu layout claro, botões azuis e barra verde)
    def apply_dark_theme(root):
        style = ttk.Style(root)
        try:
            style.theme_use('clam')
        except Exception:
            pass
        bg = "#ffffff"; surf = "#FFFFFF"; fg = "#000305"; blue = "#1f6feb"; green = "#2ea043"
        root.configure(bg=bg)
        style.configure(".", background=bg, foreground=fg)
        style.configure("TLabel", background=bg, foreground=fg)
        style.configure("Title.TLabel", font=("Segoe UI", 12, "bold"), foreground=fg, background=bg)
        style.configure("Hint.TLabel", foreground="#6b7280", background=bg, font=("Segoe UI", 9))
        style.configure("Card.TFrame", background=surf, relief="flat")
        style.configure("TButton", background=blue, foreground="white", padding=8, borderwidth=0)
        style.map("TButton", background=[("active", "#2b76ff")])
        style.configure("TEntry", fieldbackground="#ffffff", foreground=fg, insertcolor=fg)
        style.configure("Green.Horizontal.TProgressbar", troughcolor=surf, background=green)

    cfg = carregar_config()
    SIGLA = cfg["sigla"]
    ANO_PROTOCOLO = cfg.get("ano", str(DEFAULT_ANO))
    valid = cfg.get("valid", DEFAULTS_VALID.copy())
    MODELO = DEFAULT_MODELO

    PLANILHA = nome_planilha(SIGLA)
    df = carregar_planilha(PLANILHA)

    root = tk.Tk()
    root.title(f"{SIGLA} – Preencher → Planilha + PDF")
    root.geometry("720x640")
    apply_dark_theme(root)

    def on_close():
        try:
            root.quit(); root.update_idletasks(); root.destroy()
        except Exception:
            pass
        os._exit(0)
    atexit.register(lambda: None)
    root.protocol("WM_DELETE_WINDOW", on_close)

    container = ttk.Frame(root, style="Card.TFrame", padding=14)
    container.grid(sticky="nsew", padx=16, pady=16)
    root.columnconfigure(0, weight=1); root.rowconfigure(0, weight=1)

    # Cabeçalho
    ttk.Label(container, text="Formulário – Requerimento", style="Title.TLabel").grid(column=0, row=0, columnspan=3, sticky="w")
    ttk.Label(container, text="Preencha os campos abaixo. Campos marcados com * são obrigatórios.", style="Hint.TLabel").grid(column=0, row=1, columnspan=3, sticky="w", pady=(0,6))

    # Config SIGLA/ANO
    cfg_frame = ttk.Frame(container, style="Card.TFrame", padding=10)
    cfg_frame.grid(column=0, row=2, columnspan=3, sticky="we", pady=(8,12))
    cfg_frame.columnconfigure(1, weight=1)

    ttk.Label(cfg_frame, text="SIGLA do setor *").grid(column=0, row=0, sticky="e", padx=(0,8), pady=4)
    var_sigla = tk.StringVar(value=SIGLA)
    ttk.Entry(cfg_frame, textvariable=var_sigla, width=10).grid(column=1, row=0, sticky="w")

    ttk.Label(cfg_frame, text="Ano do protocolo").grid(column=2, row=0, sticky="e", padx=(16,8), pady=4)
    var_ano = tk.StringVar(value=ANO_PROTOCOLO)
    ttk.Entry(cfg_frame, textvariable=var_ano, width=10).grid(column=3, row=0, sticky="w")

    def acao_salvar_cfg():
        s = _sanitize_sigla(var_sigla.get())
        a = str(var_ano.get()).strip() or ANO_PROTOCOLO
        salvar_config(s, a, valid)  # não mexe no last_req aqui
        messagebox.showinfo("Configuração", f"Salvo: SIGLA={s}, ANO={a}\nPlanilha: {nome_planilha(s)}\n(Reabra para aplicar 100%)")
    ttk.Button(cfg_frame, text="Salvar como padrão", command=acao_salvar_cfg).grid(column=4, row=0, sticky="we", padx=(16,0))

    # Form (dois cards)
    form_left = ttk.Frame(container, style="Card.TFrame", padding=10)
    form_right = ttk.Frame(container, style="Card.TFrame", padding=10)
    form_left.grid(column=0, row=3, columnspan=2, sticky="nsew")
    form_right.grid(column=2, row=3, sticky="nsew", padx=(12,0))
    container.columnconfigure(0, weight=1); container.columnconfigure(1, weight=1); container.columnconfigure(2, weight=1)

    campos = [
        ("NOME", True),
        ("ID", True),
        ("CPF", True),
        ("CURSO", True),
        ("TURMA", True),
        ("Código da oferta", True),
        ("N chamado", False),
        ("Data", True),
        ("retorno (Previsão)", True),
    ]
    vars_map: Dict[str, Any] = {}

    # Helpers de validação
    def only_digits(s: str) -> str:
        return re.sub(r"\D", "", s or "")

    def make_numeric_limited(var: 'tk.StringVar', maxlen: int):
        def cb(*_):
            var.set(only_digits(var.get())[:maxlen])
        return cb

    def make_numeric_unlimited(var: 'tk.StringVar'):
        def cb(*_):
            var.set(only_digits(var.get()))
        return cb

    # Left: Nome, Curso, Turma, Datas
    row_l = 0
    for label in ("NOME","CURSO","TURMA","Data","retorno (Previsão)"):
        ttk.Label(form_left, text=label + (" *" if (label, True) in campos else "")).grid(column=0, row=row_l, sticky="w", pady=(2,2))
        v = tk.StringVar()
        ttk.Entry(form_left, textvariable=v, width=36).grid(column=1, row=row_l, sticky="we", padx=(8,0))
        vars_map[label] = v
        row_l += 1
    hoje = datetime.today().strftime("%d/%m/%Y")
    vars_map["Data"].set(hoje)
    vars_map["retorno (Previsão)"].set(hoje)

    # Right: CPF, ID, Código da oferta, N chamado
    row_r = 0
    ttk.Label(form_right, text="CPF *").grid(column=0, row=row_r, sticky="w", pady=(2,2))
    vcpf = tk.StringVar()
    ttk.Entry(form_right, textvariable=vcpf, width=28).grid(column=1, row=row_r, sticky="we", padx=(8,0))
    vcpf.trace_add("write", make_numeric_limited(vcpf, valid["cpf_len"]))
    vars_map["CPF"] = vcpf
    row_r += 1

    ttk.Label(form_right, text=f"ID * (máx {valid['id_len']})").grid(column=0, row=row_r, sticky="w", pady=(2,2))
    vid = tk.StringVar()
    ttk.Entry(form_right, textvariable=vid, width=28).grid(column=1, row=row_r, sticky="we", padx=(8,0))
    vid.trace_add("write", make_numeric_limited(vid, valid["id_len"]))
    vars_map["ID"] = vid
    row_r += 1

    ttk.Label(form_right, text=f"Código da oferta * (máx {valid['oferta_len']})").grid(column=0, row=row_r, sticky="w", pady=(2,2))
    vof = tk.StringVar()
    ttk.Entry(form_right, textvariable=vof, width=28).grid(column=1, row=row_r, sticky="we", padx=(8,0))
    vof.trace_add("write", make_numeric_limited(vof, valid["oferta_len"]))
    vars_map["Código da oferta"] = vof
    row_r += 1

    ttk.Label(form_right, text="N chamado").grid(column=0, row=row_r, sticky="w", pady=(2,2))
    vch = tk.StringVar()
    ttk.Entry(form_right, textvariable=vch, width=28).grid(column=1, row=row_r, sticky="we", padx=(8,0))
    vch.trace_add("write", make_numeric_unlimited(vch))
    vars_map["N chamado"] = vch
    row_r += 1

    # Barra de progresso
    ttk.Separator(container).grid(column=0, row=4, columnspan=3, sticky="we", pady=(12,8))
    progress = ttk.Progressbar(container, mode="determinate", style="Green.Horizontal.TProgressbar", maximum=100, value=0)
    progress.grid(column=0, row=5, columnspan=3, sticky="we", pady=(0,8))

    def set_progress(p):
        progress["value"] = p
        progress.update_idletasks()

    def submeter():
        nonlocal df
        set_progress(5)

        # SIGLA/ANO atuais
        sigla_atual = _sanitize_sigla(var_sigla.get())
        ano_atual = str(var_ano.get()).strip() or ANO_PROTOCOLO

        plan = nome_planilha(sigla_atual)
        if plan != nome_planilha(SIGLA):
            df_local = carregar_planilha(plan)
        else:
            df_local = df

        # valida obrigatórios
        falhas = []
        dados: Dict[str, Any] = {}
        for label, required in campos:
            val = vars_map[label].get().strip()
            if required and not val:
                falhas.append(label)
            dados[label] = val

        # validações específicas
        cpf_digits = sanitize_num(dados["CPF"])
        if len(cpf_digits) != valid["cpf_len"]:
            falhas.append(f"CPF ({valid['cpf_len']} dígitos)")
        if len(sanitize_num(dados["ID"])) != valid["id_len"]:
            falhas.append(f"ID ({valid['id_len']} dígitos)")
        if len(sanitize_num(dados["Código da oferta"])) != valid["oferta_len"]:
            falhas.append(f"Código da oferta ({valid['oferta_len']} dígitos)")

        if falhas:
            set_progress(0)
            messagebox.showwarning("Campos obrigatórios/validação", "Ajuste: " + ", ".join(sorted(set(falhas))))
            return

        set_progress(25)

        # Já existe?
        nreq_existente = existe_registro(df_local, dados)
        if nreq_existente is not None:
            try:
                set_progress(55)
                saida = gerar_documento(dados, nreq_existente, sigla_atual, ano_atual, DEFAULT_MODELO)
                set_progress(100)
                messagebox.showinfo("Já cadastrado",
                                    f"Esse registro já estava na planilha (N req. {nreq_existente}).\nArquivo em:\n{os.path.abspath(saida)}")
                set_progress(0)
            except Exception as e:
                set_progress(0)
                messagebox.showerror("Erro ao gerar documento", str(e))
            return

        # Novo registro → próximo N req. (planilha; se vazia, usa fallback do config)
        cfg_atual = carregar_config()
        nreq = proximo_nreq(df_local, cfg_atual.get("last_req", 0))
        dados["N req."] = nreq

        # Append 1 linha
        linha_df = pd.DataFrame([dados], columns=COLUNAS_PADRAO)
        df_local = pd.concat([df_local, linha_df], ignore_index=True)

        try:
            set_progress(60)
            salvar_planilha(df_local, plan)
            set_progress(75)
        except Exception as e:
            set_progress(0)
            messagebox.showerror("Erro ao salvar planilha", str(e))
            return

        # Atualiza config com o último número usado
        try:
            salvar_config(sigla_atual, ano_atual, valid, last_req=nreq)
        except Exception:
            pass

        # Gera DOCX/PDF
        try:
            saida = gerar_documento(dados, nreq, sigla_atual, ano_atual, DEFAULT_MODELO)
            set_progress(100)
        except Exception as e:
            set_progress(0)
            messagebox.showerror("Erro ao gerar documento", str(e))
            return

        if plan == nome_planilha(SIGLA):
            df = df_local

        messagebox.showinfo(
            "Pronto",
            f"Linha adicionada (N req. {nreq}).\nPlanilha: {os.path.abspath(plan)}\nArquivo em:\n{os.path.abspath(saida)}"
        )
        set_progress(0)

        # limpa campos (mantém datas)
        for label, required in campos:
            if label in ("Data","retorno (Previsão)"):
                continue
            vars_map[label].set("")

    ttk.Button(container, text="Salvar na planilha e gerar PDF", command=submeter).grid(column=0, row=6, columnspan=3, sticky="we", pady=(4,0))

    root.mainloop()

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Formulário → Planilha + PDF (multi-sigla, sem duplicar, dark+progress)")
    ap.add_argument("--gui", action="store_true", help="Abrir a interface gráfica (recomendado)")
    ap.add_argument("--sigla", type=str, help="SIGLA do setor (ex.: MCI, MMD, IOT). Se informada, sobrepõe a config para esta execução.")
    ap.add_argument("--ano", type=str, help="Ano do protocolo (ex.: 2025). Se informado, sobrepõe a config para esta execução.")
    ap.add_argument("--cpf-len", type=int, help="Tamanho do CPF (padrão 11)")
    ap.add_argument("--id-len", type=int, help="Tamanho do ID (padrão 10)")
    ap.add_argument("--oferta-len", type=int, help="Tamanho do Código da oferta (padrão 10)")
    args = ap.parse_args()

    if args.sigla or args.ano or args.cpf_len or args.id_len or args.oferta_len:
        cfg = carregar_config()
        sig = _sanitize_sigla(args.sigla or cfg["sigla"])
        ano = str(args.ano or cfg.get("ano", DEFAULT_ANO))
        valid = cfg.get("valid", DEFAULTS_VALID.copy())
        if args.cpf_len: valid["cpf_len"] = int(args.cpf_len)
        if args.id_len: valid["id_len"] = int(args.id_len)
        if args.oferta_len: valid["oferta_len"] = int(args.oferta_len)
        salvar_config(sig, ano, valid)

    iniciar_form()
