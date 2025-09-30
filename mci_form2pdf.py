#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Form → Planilha + PDF (multi-sigla, versão corrigida)
 - SIGLA configurável e persistente (MCI, MMD, IOT, etc.)
 - Cria planilha como MALA<SIGLA>.xlsx (ex.: MALAMMD.xlsx)
 - Gera DOCX/PDF em "Requerimentos/"
 - Mantém sequência do protocolo: pega o maior entre a planilha e a pasta e soma +1
 - Se (Nome+ID+CPF+Data) já existir na planilha, NÃO duplica a linha — só (re)gera o arquivo.
"""

import os
import re
import json
from datetime import datetime
from typing import Dict, Any, Optional

import pandas as pd
from docx import Document

# ---------------- Config base ----------------
DEFAULT_ANO = os.environ.get("MCI_ANO", "2025")
DEFAULT_MODELO = os.environ.get("MCI_DOCX", "anexo_geduc_se_requerimento_amparo_legall.docx")
PASTA_SAIDA = os.environ.get("MCI_SAIDAS", "Requerimentos")
CONFIG_PATH = os.environ.get("MCI_CONFIG", "config_form.json")

MESES = {1:"janeiro",2:"fevereiro",3:"março",4:"abril",5:"maio",6:"junho",7:"julho",8:"agosto",9:"setembro",10:"outubro",11:"novembro",12:"dezembro"}

COLUNAS_PADRAO = [
    "N req.","N chamado","NOME","ID","CPF","CURSO","TURMA",
    "Código da oferta","Data","retorno (Previsão)"
]

# ---------------- Utilidades de SIGLA ----------------
def _sanitize_sigla(s: str) -> str:
    s = (s or "").strip().upper()
    s = re.sub(r"[^A-Z0-9]", "", s)
    return s or "MCI"

def carregar_config() -> Dict[str, Any]:
    cfg = {"sigla": _sanitize_sigla(os.environ.get("MCI_SIGLA", "MCI")), "ano": str(DEFAULT_ANO)}
    if os.path.isfile(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                if "sigla" in data:
                    cfg["sigla"] = _sanitize_sigla(data["sigla"])
                if "ano" in data and str(data["ano"]).isdigit():
                    cfg["ano"] = str(data["ano"])
        except Exception:
            pass
    return cfg

def salvar_config(sigla: str, ano: Optional[str] = None) -> None:
    data = {"sigla": _sanitize_sigla(sigla), "ano": str(ano or DEFAULT_ANO)}
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def nome_planilha(sigla: str) -> str:
    # Se quiser forçar um caminho via env, ainda é possível.
    forced = os.environ.get("MCI_XLSX", "").strip()
    return forced if forced else f"MALA{_sanitize_sigla(sigla)}.xlsx"

# ---------------- Funções auxiliares ----------------
def fmt_data_pt(dt: datetime) -> str:
    return f"{dt.day} de {MESES[dt.month]} de {dt.year}"

def sanitize_num(s: Any) -> str:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    return re.sub(r"\D", "", str(s))

def get_str(s: Any) -> str:
    return "" if (s is None or (isinstance(s, float) and pd.isna(s))) else str(s).strip()

def parse_data_flex(valor: str, fallback_hoje: bool = True) -> datetime:
    if not valor:
        return datetime.today() if fallback_hoje else None
    for fmt in ("%d/%m/%Y","%Y-%m-%d","%d-%m-%Y","%d.%m.%Y"):
        try:
            return datetime.strptime(str(valor), fmt)
        except Exception:
            pass
    try:
        dt = pd.to_datetime(valor, dayfirst=True, errors="coerce")
        if pd.isna(dt):
            raise ValueError
        return dt.to_pydatetime()
    except Exception:
        return datetime.today() if fallback_hoje else None

def carregar_planilha(caminho: str) -> pd.DataFrame:
    if os.path.isfile(caminho):
        df = pd.read_excel(caminho)
        df.columns = [c.strip() for c in df.columns]
        return df
    return pd.DataFrame(columns=COLUNAS_PADRAO)

def salvar_planilha(df: pd.DataFrame, caminho: str) -> None:
    for c in COLUNAS_PADRAO:
        if c not in df.columns:
            df[c] = None
    df.to_excel(caminho, index=False)

def max_protocolo_na_pasta(sigla: str, ano: str) -> int:
    """
    Varre "Requerimentos" e acha o maior NN em 'NN<SIGLA>{ANO} ...'.
    """
    if not os.path.isdir(PASTA_SAIDA):
        return 0
    maxp = 0
    sigla = _sanitize_sigla(sigla)
    padrao = re.compile(rf'^(\d{{2}}){re.escape(sigla)}{re.escape(ano)}\b', re.IGNORECASE)
    for nome in os.listdir(PASTA_SAIDA):
        m = padrao.match(nome)
        if m:
            try:
                maxp = max(maxp, int(m.group(1)))
            except Exception:
                pass
    return maxp

def proximo_nreq(df: pd.DataFrame, sigla: str, ano: str) -> int:
    max_sheet = 0
    if "N req." in df.columns and not df.empty:
        nums = pd.to_numeric(df["N req."], errors="coerce")
        if not nums.dropna().empty:
            max_sheet = int(nums.max())
    max_folder = max_protocolo_na_pasta(sigla, ano)
    return max(max_sheet, max_folder) + 1

def existe_registro(df: pd.DataFrame, dados: Dict[str, Any]) -> Optional[int]:
    if df.empty:
        return None
    filtros = (df.get("NOME","").astype(str).str.strip().str.casefold() == str(dados.get("NOME","")).strip().casefold())
    for col in ("ID","CPF","Data"):
        if col in df.columns:
            filtros = filtros & (df[col].astype(str).str.strip() == str(dados.get(col,"")).strip())
    existentes = df[filtros]
    if existentes.empty:
        return None
    try:
        return int(pd.to_numeric(existentes.iloc[0]["N req."], errors="coerce"))
    except Exception:
        return None

def montar_mapa(row: Dict[str, Any], protocolo_num: int) -> Dict[str, str]:
    dt_solic = parse_data_flex(row.get("Data"))
    dt_retorno = parse_data_flex(row.get("retorno (Previsão)"))
    return {
        "protocolo": f"{int(protocolo_num):02d}",
        "nome": get_str(row.get("NOME")),
        "id": sanitize_num(row.get("ID")),
        "cpf": sanitize_num(row.get("CPF")),
        "curso": get_str(row.get("CURSO")),
        "turma": get_str(row.get("TURMA")),
        "oferta": sanitize_num(row.get("Código da oferta")),
        "chamado": sanitize_num(row.get("N chamado")),
        "data_ext": fmt_data_pt(dt_solic),
        "data_retorno_ext": fmt_data_pt(dt_retorno),
    }

def substituir_texto(doc: Document, mapa: Dict[str, str], sigla: str, ano: str) -> None:
    sigla = _sanitize_sigla(sigla)
    pares = [
        (rf"Protocolo nº \d{{2}}-{sigla}/\d{{4}}", f"Protocolo nº {mapa['protocolo']}-{sigla}/{ano}"),
        (r"Eu, .+?, ID nº .*?; CPF nº .*?, estudante regularmente matriculado\(a\) no curso .*?,\s*TURMA:.*?,",
         f"Eu, {mapa['nome']}, ID nº {mapa['id']}; CPF nº {mapa['cpf']}, estudante regularmente matriculado(a) no curso {mapa['curso']}, TURMA:{mapa['turma']},"),
        (r"Código da oferta: .*? \(preenchimento do setor da secretaria escolar\)",
         f"Código da oferta: {mapa['oferta']} (preenchimento do setor da secretaria escolar)"),
        (r"São Paulo, .*?\.", f"São Paulo, {mapa['data_ext']}."),
        (r"Conforme chamado de nº .*", f"Conforme chamado de nº {mapa['chamado']}"),
        (r"Aluno .+", f"Aluno {mapa['nome']}"),
        (r"Data de retorno até: .*?\s+\(considerar.*\)",
         f"Data de retorno até: {mapa['data_retorno_ext']}  (considerar de 1 a 7 dias úteis, a partir da data de solicitação)"),
        # caso o modelo esteja com outra sigla antiga, normalizamos qualquer sigla:
        (r"Protocolo nº \d{2}-[A-Z0-9]+/\d{4}", f"Protocolo nº {mapa['protocolo']}-{sigla}/{ano}"),
    ]

    def apply_replace(text: str) -> str:
        for pattern, repl in pares:
            text = re.sub(pattern, repl, text)
        return text

    # Parágrafos
    for p in doc.paragraphs:
        if p.text:
            joined = "".join(run.text for run in p.runs)
            new = apply_replace(joined)
            if new != joined:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""

    # Tabelas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        joined = "".join(run.text for run in p.runs)
                        new = apply_replace(joined)
                        if new != joined:
                            p.runs[0].text = new
                            for r in p.runs[1:]:
                                r.text = ""

def gerar_documento(linha: Dict[str, Any], protocolo_num: int, sigla: str, ano: str, modelo: str) -> str:
    os.makedirs(PASTA_SAIDA, exist_ok=True)
    mapa = montar_mapa(linha, protocolo_num)
    base = f"{mapa['protocolo']}{_sanitize_sigla(sigla)}{ano} {mapa['nome']}"
    out_docx = os.path.join(PASTA_SAIDA, base + ".docx")
    out_pdf = os.path.join(PASTA_SAIDA, base + ".pdf")

    # evita duplicar se já existir
    if os.path.isfile(out_pdf):
        return out_pdf
    if os.path.isfile(out_docx):
        # tenta converter para pdf se possível
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(out_docx, out_pdf)
            return out_pdf
        except Exception:
            return out_docx

    doc = Document(modelo)
    substituir_texto(doc, mapa, sigla, ano)
    doc.save(out_docx)

    saida_final = out_docx
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(out_docx, out_pdf)
        saida_final = out_pdf
    except Exception:
        pass
    return saida_final

# ---------------- GUI ----------------
def iniciar_form():
    import tkinter as tk
    from tkinter import ttk, messagebox

    cfg = carregar_config()
    SIGLA = cfg["sigla"]
    ANO_PROTOCOLO = cfg.get("ano", str(DEFAULT_ANO))
    MODELO = DEFAULT_MODELO

    PLANILHA = nome_planilha(SIGLA)
    df = carregar_planilha(PLANILHA)

    root = tk.Tk()
    root.title(f"{SIGLA} – Preencher → Planilha + PDF")
    frm = ttk.Frame(root, padding=12)
    frm.grid(sticky="nsew")
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)

    # ---- Linha de config SIGLA/ANO ----
    r = 0
    ttk.Label(frm, text="Configuração").grid(column=0, row=r, columnspan=3, sticky="w")
    r += 1

    ttk.Label(frm, text="SIGLA do setor *").grid(column=0, row=r, sticky="e", padx=(0,8), pady=4)
    var_sigla = tk.StringVar(value=SIGLA)
    ent_sigla = ttk.Entry(frm, textvariable=var_sigla, width=12)
    ent_sigla.grid(column=1, row=r, sticky="w")

    ttk.Label(frm, text="Ano do protocolo").grid(column=0, row=r+1, sticky="e", padx=(0,8), pady=4)
    var_ano = tk.StringVar(value=ANO_PROTOCOLO)
    ent_ano = ttk.Entry(frm, textvariable=var_ano, width=12)
    ent_ano.grid(column=1, row=r+1, sticky="w")

    def acao_salvar_cfg():
        s = _sanitize_sigla(var_sigla.get())
        a = str(var_ano.get()).strip() or ANO_PROTOCOLO
        salvar_config(s, a)
        messagebox.showinfo("Configuração", f"Salvo: SIGLA={s}, ANO={a}\n"
                            f"Planilha passará a ser: {nome_planilha(s)}\n"
                            f"(Reabra a janela para aplicar totalmente)")
    ttk.Button(frm, text="Salvar como padrão", command=acao_salvar_cfg).grid(column=2, row=r, rowspan=2, sticky="we", padx=8)

    r += 2
    ttk.Separator(frm, orient="horizontal").grid(column=0, row=r, columnspan=3, sticky="we", pady=(8,8))
    r += 1

    # ---- Campos do formulário ----
    campos = [
        ("NOME", True),
        ("ID", True),
        ("CPF", True),
        ("CURSO", True),
        ("TURMA", True),
        ("Código da oferta", True),
        ("N chamado", False),
        ("Data", True),
        ("retorno (Previsão)", True),
    ]

    vars_map = {}
    ttk.Label(frm, text="Preencha os dados do aluno").grid(column=0, row=r, columnspan=3, sticky="w")
    r += 1

    for label, required in campos:
        ttk.Label(frm, text=label + ("" if not required else " *")).grid(column=0, row=r, sticky="e", padx=(0,8), pady=4)
        v = tk.StringVar()
        ent = ttk.Entry(frm, textvariable=v, width=50)
        ent.grid(column=1, row=r, columnspan=2, sticky="we")
        vars_map[label] = v
        r += 1

    hoje = datetime.today().strftime("%d/%m/%Y")
    vars_map["Data"].set(hoje)
    vars_map["retorno (Previsão)"].set(hoje)

    def submeter():
        nonlocal df
        # Atualiza SIGLA/ANO em runtime (sem precisar reabrir)
        sigla_atual = _sanitize_sigla(var_sigla.get())
        ano_atual = str(var_ano.get()).strip() or ANO_PROTOCOLO

        plan = nome_planilha(sigla_atual)
        if plan != nome_planilha(SIGLA):
            # se o usuário trocou a sigla agora, recarrega/abre a nova planilha
            df_local = carregar_planilha(plan)
        else:
            df_local = df

        # valida
        falhas = []
        dados = {}
        for label, required in campos:
            val = vars_map[label].get().strip()
            if required and not val:
                falhas.append(label)
            dados[label] = val
        if falhas:
            messagebox.showwarning("Campos obrigatórios", "Preencha: " + ", ".join(falhas))
            return

        # Já existe na planilha?
        nreq_existente = existe_registro(df_local, dados)
        if nreq_existente is not None:
            try:
                saida = gerar_documento(dados, nreq_existente, sigla_atual, ano_atual, MODELO)
                messagebox.showinfo("Já cadastrado",
                                    f"Esse registro já estava na planilha (N req. {nreq_existente}).\nArquivo em:\n{os.path.abspath(saida)}")
            except Exception as e:
                messagebox.showerror("Erro ao gerar documento", str(e))
            return

        # Novo registro → próximo N req.
        nreq = proximo_nreq(df_local, sigla_atual, ano_atual)
        dados["N req."] = nreq

        # Append 1 linha
        linha_df = pd.DataFrame([dados], columns=COLUNAS_PADRAO)
        df_local = pd.concat([df_local, linha_df], ignore_index=True)
        try:
            salvar_planilha(df_local, plan)
        except Exception as e:
            messagebox.showerror("Erro ao salvar planilha", str(e))
            return

        try:
            saida = gerar_documento(dados, nreq, sigla_atual, ano_atual, MODELO)
        except Exception as e:
            messagebox.showerror("Erro ao gerar documento", str(e))
            return

        # atualiza referência em memória se a planilha era a mesma
        if plan == nome_planilha(SIGLA):
            df = df_local

        messagebox.showinfo("Pronto",
                            f"Linha adicionada (N req. {nreq}).\nPlanilha: {os.path.abspath(plan)}\nArquivo em:\n{os.path.abspath(saida)}")

        # limpa campos (mantém datas)
        for label, required in campos:
            if label in ("Data","retorno (Previsão)"):
                continue
            vars_map[label].set("")

    ttk.Button(frm, text="Salvar na planilha e gerar PDF", command=submeter).grid(column=0, row=r, columnspan=3, sticky="we", pady=12)
    ttk.Label(frm, text="* Campos obrigatórios").grid(column=0, row=r+1, columnspan=3, sticky="w")

    root.mainloop()

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Formulário → Planilha + PDF (multi-sigla, sem duplicar)")
    ap.add_argument("--gui", action="store_true", help="Abrir a interface gráfica (recomendado)")
    ap.add_argument("--sigla", type=str, help="SIGLA do setor (ex.: MCI, MMD, IOT). Se informada, sobrepõe a config para esta execução.")
    ap.add_argument("--ano", type=str, help="Ano do protocolo (ex.: 2025). Se informado, sobrepõe a config para esta execução.")
    args = ap.parse_args()

    if args.gui or True:
        # Se o usuário passou overrides por CLI, salva/usa agora (efeito imediato e persistente)
        if args.sigla or args.ano:
            cfg = carregar_config()
            sig = _sanitize_sigla(args.sigla or cfg["sigla"])
            ano = str(args.ano or cfg.get("ano", DEFAULT_ANO))
            salvar_config(sig, ano)
        iniciar_form()
